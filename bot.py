import os
import logging
from aiogram import Bot, Dispatcher, types
from aiogram.contrib.middlewares.logging import LoggingMiddleware
from aiogram.contrib.fsm_storage.memory import MemoryStorage
from aiogram.dispatcher import FSMContext
from aiogram.dispatcher.filters.state import State, StatesGroup
from aiogram.utils import executor
from docx import Document
from io import BytesIO

# Получение токена из переменной окружения
API_TOKEN = os.getenv('API_TOKEN')

# Настройка логирования
logging.basicConfig(level=logging.INFO)

# Инициализация бота и диспетчера
bot = Bot(token=API_TOKEN)
storage = MemoryStorage()
dp = Dispatcher(bot, storage=storage)
dp.middleware.setup(LoggingMiddleware())

# Определение состояний
class Form(StatesGroup):
    name1 = State()
    name2 = State()
    dob = State()
    address = State()
    email = State()
    passport_number = State()
    passport_issuer = State()
    snils = State()
    phone = State()
    direction = State()

# Обработчик команды /start
@dp.message_handler(commands='start')
async def cmd_start(message: types.Message):
    logging.info("Команда /start получена")
    await Form.name1.set()
    await message.reply("Привет! Я помогу тебе создать заявление. Введи свои фамилию, имя и отчество в именительном падеже (напр. Иванов Иван Иванович).")

# Обработчик ввода имени и фамилии
@dp.message_handler(state=Form.name1)
async def process_name(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['name1'] = message.text
    await Form.next()
    await message.reply("Теперь введи их в родительном падеже (напр. Иванова Ивана Ивановича).")
    
# Обработчик ввода имени и фамилии
@dp.message_handler(state=Form.name2)
async def process_name(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['name2'] = message.text
    await Form.next()
    await message.reply("Введите вашу дату рождения (ДД.ММ.ГГГГ).")

# Обработчик ввода даты рождения
@dp.message_handler(state=Form.dob)
async def process_dob(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['dob'] = message.text
    await Form.next()
    await message.reply("Введите ваш адрес регистрации.")

# Обработчик ввода адреса
@dp.message_handler(state=Form.address)
async def process_address(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['address'] = message.text
    await Form.next()
    await message.reply("Введите ваш email.")

# Обработчик ввода email
@dp.message_handler(state=Form.email)
async def process_email(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['email'] = message.text
    await Form.next()
    await message.reply("Введите серию и номер паспорта, отделите их пробелом.")

# Обработчик ввода серии и номера паспорта
@dp.message_handler(state=Form.passport_number)
async def process_passport_number(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['passport_number'] = message.text
    await Form.next()
    await message.reply("Кем выдан паспорт?")

# Обработчик ввода данных о выдаче паспорта
@dp.message_handler(state=Form.passport_issuer)
async def process_passport_issuer(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['passport_issuer'] = message.text
    await Form.next()
    await message.reply("Введите ваш СНИЛС.")

# Обработчик ввода СНИЛС
@dp.message_handler(state=Form.snils)
async def process_snils(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['snils'] = message.text
    await Form.next()
    await message.reply("Введите ваш контактный телефон в виде +7-ХХХ-ХХХ-ХХ-ХХ.")

# Обработчик ввода контактного телефона
@dp.message_handler(state=Form.phone)
async def process_phone(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['phone'] = message.text
    await Form.next()
    await message.reply("Выберите направление:\n1. Анализ данных в педагогической деятельности\n2. Автоматизация работы с данными и документами")

# Обработчик выбора направления
@dp.message_handler(state=Form.direction)
async def process_direction(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        if message.text == "1":
            data['direction'] = "Анализ данных в педагогической деятельности"
        elif message.text == "2":
            data['direction'] = "Автоматизация работы с данными и документами"
        else:
            await message.reply("Пожалуйста, выберите корректное направление: 1 или 2")
            return

        # Загрузка шаблона документа
        template_path = 'template.docx'
        doc = Document(template_path)

        # Замена заполнителей в документе
        replace_text(doc, '{{name1}}', data['name1'])
        replace_text(doc, '{{name2}}', data['name2'])
        replace_text(doc, '{{dob}}', data['dob'])
        replace_text(doc, '{{address}}', data['address'])
        replace_text(doc, '{{email}}', data['email'])
        replace_text(doc, '{{passport_number}}', data['passport_number'])
        replace_text(doc, '{{passport_issuer}}', data['passport_issuer'])
        replace_text(doc, '{{snils}}', data['snils'])
        replace_text(doc, '{{phone}}', data['phone'])
        replace_text(doc, '{{direction}}', data['direction'])

        # Сохранение документа в BytesIO
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        input_file = types.InputFile(file_stream, filename='zayavlenie.docx')

        # Отправка документа пользователю
        await bot.send_document(chat_id=message.chat.id, document=input_file)
        await message.reply("Заявление создано и отправлено!")

    # Завершение состояния
    await state.finish()

# Функция для замены текста в документе
def replace_text(doc, placeholder, replacement):
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, replacement)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text(cell, placeholder, replacement)

if __name__ == '__main__':
    logging.info("Запуск бота")
    executor.start_polling(dp, skip_updates=True)
